"""
ingestion.py — Document Ingestion Pipeline
Handles loading, parsing, and chunking of PDF and DOCX files.
"""
import os
from typing import List, Tuple
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from config import CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS


def load_pdf(file_path: str) -> str:
    try:
        pages = PyPDFLoader(file_path).load()
        text = "\n".join([p.page_content for p in pages])
        if text.strip():
            return text
    except Exception:
        pass
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def load_docx(file_path: str) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def load_document(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported format: '{ext}'. Supported: {SUPPORTED_EXTENSIONS}")


def chunk_text(text: str, source: str, chunk_size=None, chunk_overlap=None) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or CHUNK_SIZE,
        chunk_overlap=chunk_overlap or CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    docs = [Document(page_content=text, metadata={"source": source})]
    chunks = splitter.split_documents(docs)
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        chunk.metadata["total_chunks"] = len(chunks)
    return chunks


def ingest_document(file_path: str, filename: str, chunk_size=None, chunk_overlap=None) -> Tuple[List[Document], str]:
    text = load_document(file_path)
    if not text.strip():
        raise ValueError(f"Document '{filename}' is empty.")
    chunks = chunk_text(text, filename, chunk_size, chunk_overlap)
    print(f"[Ingestion] '{filename}' -> {len(chunks)} chunks")
    return chunks, text
